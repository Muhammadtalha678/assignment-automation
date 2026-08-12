import json
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from src.helper.helper_functions import add_page_border,set_table_borders
from src.controllers.graphviz_diagram import generate_graphviz_diagram
from src.controllers.generate_image import generate_image, generate_image_via_advanced_web

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Utility to set internal cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

async def generate_assignment_docx(image_map:str,json_data_str: str, output_path: str = "Assignment.docx", logo_path: str = None):
    # 1. Parse JSON Input
    if isinstance(json_data_str, str):
        data = json.loads(json_data_str)
    else:
        data = json_data_str

    # assignment_info = data.get("assignment", {})
    questions = data.get("questions", [])

    # -------------------------------------------------------------------------
    # PHASE 1:Already GENERATED ALL IMAGES
    # -------------------------------------------------------------------------
    # 2. Initialize Document
    doc = Document()
    add_page_border(doc)

    # Set Document Page Margins (Normal 1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # -------------------------------------------------------------
    # HEADER / LOGO SECTION
    # -------------------------------------------------------------
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.8))
        p_logo.paragraph_format.space_after = Pt(12)

    # Main Header: Assignment # 01 (Calibri)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"Assignment # 0{data.get('assignment_no', 1)}")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66) # Deep Navy
    p_title.paragraph_format.space_after = Pt(18)

    # -------------------------------------------------------------
    # STUDENT DETAILS TABLE
    # -------------------------------------------------------------
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # table.autofit = False
    details = [
        ("Student Name:", str(data.get("student_name", ""))),
        ("Registration ID:", f"0000{data.get('registration_id', '')}"),
        ("Course Code:", str(data.get("course_code", ""))),
        ("Semester:", str(data.get("semester", "")))
    ]

    for i, (label, val) in enumerate(details):
        row = table.rows[i]
        
        # Label Cell
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(4)
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Calibri'
        run_lbl.font.size = Pt(11)
        run_lbl.font.bold = True

        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_lbl.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Value Cell
        cell_val = row.cells[1]
        cell_val.width = Inches(4.3)
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        run_val = p_val.add_run(val)
        run_val.font.name = 'Calibri'
        run_val.font.size = Pt(11)

        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)

    # Add light border/shading to table cells
    for row in table.rows:
        for cell in row.cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # -------------------------------------------------------------
    # QUESTIONS & ANSWERS LOOP
    # -------------------------------------------------------------
    for q_idx, q in enumerate(questions, start=1):
        # if q_idx > 1:
            # doc.add_page_break()
        # --- Question Heading (Times New Roman, Size 16, Bold) ---
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(18)
        p_q.paragraph_format.space_after = Pt(8)
        p_q.paragraph_format.keep_with_next = True
        if q_idx > 1:
            p_q.paragraph_format.page_break_before = True
        run_q = p_q.add_run(f"Q.{q.get('question_number', q_idx)}: {q.get('question_text', '')}")
        run_q.font.name = 'Times New Roman'
        run_q.font.size = Pt(16)
        run_q.font.bold = True
        run_q.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

        # --- Answer Introduction ---
        if q.get("introduction"):
            intro_h = doc.add_paragraph()
            intro_h.paragraph_format.space_before = Pt(12)
            intro_h.paragraph_format.space_after = Pt(4)
            intro_h.paragraph_format.keep_with_next = True
            
            run_intro_h = intro_h.add_run("Introduction")
            run_intro_h.font.name = 'Times New Roman'
            run_intro_h.font.size = Pt(14)
            run_intro_h.font.bold = True
            run_intro_h.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

            # Text ko new lines (\n) par split karein aur alag paragraphs banayein
            intro_text = q.get("introduction", "")
            paragraphs_list = [p.strip() for p in intro_text.replace("\\n", "\n").split("\n") if p.strip()]
            for para_text in paragraphs_list:
                p_intro = doc.add_paragraph()
                p_intro.paragraph_format.space_after = Pt(10)
                p_intro.paragraph_format.line_spacing = 1.15
                run_intro = p_intro.add_run(para_text)
                run_intro.font.name = 'Times New Roman'
                run_intro.font.size = Pt(12)

        # --- Sections / Sub-headings (Times New Roman, Size 14, Bold) ---
        for sec in q.get("sections", []):
            p_h = doc.add_paragraph()
            p_h.paragraph_format.space_before = Pt(12)
            p_h.paragraph_format.space_after = Pt(4)
            p_h.paragraph_format.keep_with_next = True

            run_h = p_h.add_run(sec.get("heading", ""))
            run_h.font.name = 'Times New Roman'
            run_h.font.size = Pt(14)
            run_h.font.bold = True
            run_h.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

             # Text ko new lines (\n) par split karein aur alag paragraphs banayein
            body_text = sec.get("explanation", "")
            paragraphs_list = [p.strip() for p in body_text.replace("\\n", "\n").split("\n") if p.strip()]
            for para_text in paragraphs_list:
                p_body = doc.add_paragraph()
                p_body.paragraph_format.space_after = Pt(10)
                p_body.paragraph_format.line_spacing = 1.15
                run_body = p_body.add_run(para_text)
                run_body.font.name = 'Times New Roman'
                run_body.font.size = Pt(12)

        # --- Diagram Description Box ---
        # Add Pre-generated Image from Phase 1
        img_path = image_map.get(q_idx)
      
        #     doc.add_paragraph().paragraph_format.space_after = Pt(10)
        if img_path and os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)   # Small spacing before image
            p_img.paragraph_format.space_after = Pt(6)    # Small spacing after image

            p_img.paragraph_format.keep_with_next = False
            run_img = p_img.add_run()
            
            # # Set maximum width AND restrict maximum height so it fits seamlessly on the page
            # # Height <= 2.2 inches ensure karega ke question ke baad white space bilkul na bachay
            
            run_img.add_picture(img_path, width=Inches(5.8))

            
            # # Caption
            # p_cap = doc.add_paragraph()
            # p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # run_cap = p_cap.add_run(f"Figure {idx}: {q.get('diagram_description')}")
            # run_cap.font.name = 'Times New Roman'
            # run_cap.font.size = Pt(10)
            # run_cap.font.italic = True
            # run_cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # --- Conclusion Section (Times New Roman, Size 14, Bold) ---
        if q.get("conclusion"):
            p_conc_h = doc.add_paragraph()
            p_conc_h.paragraph_format.space_before = Pt(12)
            p_conc_h.paragraph_format.space_after = Pt(4)
            p_conc_h.paragraph_format.keep_with_next = True

            run_conc_h = p_conc_h.add_run("Conclusion")
            run_conc_h.font.name = 'Times New Roman'
            run_conc_h.font.size = Pt(14)
            run_conc_h.font.bold = True
            run_conc_h.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

            conclusion_text = q.get("conclusion", "")
            paragraphs_list = [p.strip() for p in conclusion_text.replace("\\n", "\n").split("\n") if p.strip()]
            for para_text in paragraphs_list:
                p_conc_body = doc.add_paragraph()
                p_conc_body.paragraph_format.space_after = Pt(16)
                p_conc_body.paragraph_format.line_spacing = 1.15
                run_conc_body = p_conc_body.add_run(para_text)
                run_conc_body.font.name = 'Times New Roman'
                run_conc_body.font.size = Pt(12)

    # Save Word Document
    doc.save(output_path)
    print("--> Phase 3: Cleaning up temporary images...")
    
    return output_path