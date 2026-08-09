import json
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def enable_document_rtl(doc):
    """
    Sets the entire document section to Right-To-Left mode.
    """
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

def set_urdu_paragraph(paragraph, justify=True):
    """
    Forces Right-to-Left formatting and alignment on paragraph level.
    """
    pPr = paragraph._p.get_or_add_pPr()
    
    # BiDi element for Paragraph
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

    # Justify or Right Align
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both" if justify else "right")
    pPr.append(jc)

# def set_urdu_run(run, font_name="Calibri", size_pt=13, bold=False, color_rgb=None):
#     """
#     Sets Formatting, Complex Script (CS) Font, Size, Color and RTL properties for Urdu runs.
#     """
#     # 1. Base formatting
#     run.font.bold = bold
#     run.font.name = font_name
    
#     rPr = run._r.get_or_add_rPr()
    
#     # 2. Enable RTL on run
#     rtl = OxmlElement("w:rtl")
#     rtl.set(qn("w:val"), "1")
#     rPr.append(rtl)

#     # 3. Set Fonts (ASCII, HAnsi, and Complex Script)
#     rFonts = OxmlElement("w:rFonts")
#     rFonts.set(qn("w:ascii"), font_name)
#     rFonts.set(qn("w:hAnsi"), font_name)
#     rFonts.set(qn("w:cs"), font_name)  # Complex script font (Urdu ke liye zaroori hai)
#     rPr.append(rFonts)

#     # 4. Fix Font Size for Complex Script (Urdu)
#     # Word use half-points for font size, so 14pt = 28 half-points
#     size_half_pts = str(int(size_pt * 2))
    
#     sz = OxmlElement("w:sz")
#     sz.set(qn("w:val"), size_half_pts)
#     rPr.append(sz)
    
#     szCs = OxmlElement("w:szCs")  # Complex Script Font Size fix
#     szCs.set(qn("w:val"), size_half_pts)
#     rPr.append(szCs)

#     # 5. Fix Color for Complex Script (Urdu)
#     if color_rgb:
#         # Convert RGBColor object to Hex string (e.g., 003366)
#         hex_color = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
        
#         color = OxmlElement("w:color")
#         color.set(qn("w:val"), hex_color)
#         rPr.append(color)

def set_urdu_run(run, font_name="Calibri", size_pt=13, bold=False, color_rgb=None):
    """
    Sets Formatting, Complex Script (CS) Font, Size, Color, RTL, and REAL Dark Bold properties for Urdu runs.
    """
    run.font.name = font_name
    
    rPr = run._r.get_or_add_rPr()
    
    # 1. Enable RTL on run
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)

    # 2. Set Fonts (ASCII, HAnsi, and Complex Script)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.append(rFonts)

    # 3. REAL Dark Bold Fix for Complex Script (Urdu)
    if bold:
        run.font.bold = True  # Standard bold
        
        bCs = OxmlElement("w:bCs")  # Complex Script Bold Tag (This makes it DARK bold)
        bCs.set(qn("w:val"), "1")
        rPr.append(bCs)
        
        b = OxmlElement("w:b")  # Standard Bold Tag for safety
        b.set(qn("w:val"), "1")
        rPr.append(b)

    # 4. Fix Font Size for Complex Script (Urdu)
    size_half_pts = str(int(size_pt * 2))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), size_half_pts)
    rPr.append(sz)
    
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), size_half_pts)
    rPr.append(szCs)

    # 5. Fix Color for Complex Script (Urdu)
    # if color_rgb:
    #     hex_color = f"{color_rgb.rgb[0]:02X}{color_rgb.rgb[1]:02X}{color_rgb.rgb[2]:02X}"
    #     color = OxmlElement("w:color")
    #     color.set(qn("w:val"), hex_color)
    #     rPr.append(color)
        
    if color_rgb:
#         # Convert RGBColor object to Hex string (e.g., 003366)
        hex_color = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
        
        color = OxmlElement("w:color")
        color.set(qn("w:val"), hex_color)
        rPr.append(color)

# def set_urdu_run(run, font_name="Calibri", size_pt=13, bold=False, color_rgb=None):
#     """
#     Sets Complex Script (CS) Font and RTL properties for Urdu runs.
#     Fallback font is set to 'Calibri' if Nastaleeq isn't available.
#     """
#     run.font.size = Pt(size_pt)
#     run.font.bold = bold
#     run.font.name = font_name
#     if color_rgb:
#         run.font.color.rgb = color_rgb

#     rPr = run._r.get_or_add_rPr()
    
#     # Enable RTL on run
#     rtl = OxmlElement("w:rtl")
#     rtl.set(qn("w:val"), "1")
#     rPr.append(rtl)

#     # Set ASCII, CS, and High-ANSI font attributes
#     rFonts = OxmlElement("w:rFonts")
#     rFonts.set(qn("w:ascii"), font_name)
#     rFonts.set(qn("w:hAnsi"), font_name)
#     rFonts.set(qn("w:cs"), font_name)
#     rPr.append(rFonts)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Utility to set internal cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_border(doc):
    section = doc.sections[0]
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")

    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "16")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        pgBorders.append(border)
    sectPr.append(pgBorders)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        borders.append(border)
    tblPr.append(borders)

async def generate_assignment_docx(image_map: dict, json_data_str: str, output_path: str = "Assignment.docx", logo_path: str = None):
    data = json.loads(json_data_str) if isinstance(json_data_str, str) else json_data_str
    questions = data.get("questions", [])

    doc = Document()
    
    # 1. Enable Global Document Level RTL
    enable_document_rtl(doc)
    add_page_border(doc)

    # 2. Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 3. Header / Logo
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(10)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.5))

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"Assignment # 0{data.get('assignment_no', 1)}")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    p_title.paragraph_format.space_after = Pt(14)

    # 4. Student Details Table
    table = doc.add_table(rows=4, cols=2)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    details = [
        ("Student Name:", str(data.get("student_name", ""))),
        ("Registration ID:", f"0000{data.get('registration_id', '')}"),
        ("Course Code:", str(data.get("course_code", ""))),
        ("Semester:", str(data.get("semester", "")))
    ]

    for i, (label, val) in enumerate(details):
        row = table.rows[i]
        
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Calibri'
        run_lbl.font.size = Pt(11)
        run_lbl.font.bold = True
        cell_lbl.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        cell_val = row.cells[1]
        cell_val.width = Inches(4.3)
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p_val.add_run(val)
        run_val.font.name = 'Calibri'
        run_val.font.size = Pt(11)
        cell_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)

        for cell in (cell_lbl, cell_val):
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 5. Questions and Answers Loop
    for q_idx, q in enumerate(questions, start=1):

        if q_idx > 1:
            doc.add_page_break()
        # Question Heading
        p_q = doc.add_paragraph()
        
        set_urdu_paragraph(p_q, justify=True)
        p_q.paragraph_format.space_before = Pt(16)
        p_q.paragraph_format.space_after = Pt(8)
        p_q.paragraph_format.keep_with_next = True
        
        run_q = p_q.add_run(f"سوال نمبر {q.get('question_number')}: {q.get('question_text')}")
        set_urdu_run(run_q, font_name='Calibri', size_pt=16, bold=True,)

        # Introduction
        if q.get("introduction"):
            p_intro_h = doc.add_paragraph()
            set_urdu_paragraph(p_intro_h, justify=True)
            p_intro_h.paragraph_format.space_before = Pt(10)
            p_intro_h.paragraph_format.space_after = Pt(4)
            p_intro_h.paragraph_format.keep_with_next = True
            
            run_intro_h = p_intro_h.add_run("تعارف")
            set_urdu_run(run_intro_h, font_name="Calibri", size_pt=15, color_rgb=RGBColor(0x00, 0x33, 0x66), bold=True)
            
            # Text ko new lines (\n) par split karein aur alag paragraphs banayein
            intro_text = q.get("introduction", "")
            paragraphs_list = [p.strip() for p in intro_text.replace("\\n", "\n").split("\n") if p.strip()]
            
            for para_text in paragraphs_list:
                p_intro = doc.add_paragraph()
                set_urdu_paragraph(p_intro, justify=True)
                p_intro.paragraph_format.space_after = Pt(10)
                p_intro.paragraph_format.line_spacing = 1.35
                run_intro = p_intro.add_run(para_text)
                set_urdu_run(run_intro, font_name="Calibri", size_pt=14, bold=False)

            # p_intro = doc.add_paragraph()
            # set_urdu_paragraph(p_intro, justify=True)
            # p_intro.paragraph_format.space_after = Pt(10)
            # p_intro.paragraph_format.line_spacing = 1.35
            # run_intro = p_intro.add_run(q.get("introduction"))
            # set_urdu_run(run_intro, font_name="Calibri", size_pt=14, bold=False)

        # Sub-sections
        for sec in q.get("sections", []):
            p_h = doc.add_paragraph()
            set_urdu_paragraph(p_h, justify=True)
            p_h.paragraph_format.space_before = Pt(12)
            p_h.paragraph_format.space_after = Pt(4)
            p_h.paragraph_format.keep_with_next = True

            run_h = p_h.add_run(sec.get("heading", ""))
            set_urdu_run(run_h, font_name='Calibri', size_pt=15, bold=True, color_rgb=RGBColor(0x00, 0x33, 0x66))

             # Text ko new lines (\n) par split karein aur alag paragraphs banayein
            body_text = sec.get("explanation", "")
            paragraphs_list = [p.strip() for p in body_text.replace("\\n", "\n").split("\n") if p.strip()]

            for para_text in paragraphs_list:
                p_body = doc.add_paragraph()
                set_urdu_paragraph(p_body, justify=True)
                p_body.paragraph_format.space_after = Pt(10)
                p_body.paragraph_format.line_spacing = 1.35
                run_body = p_body.add_run(para_text)
                set_urdu_run(run_body, font_name='Calibri', size_pt=14, bold=False)

            # p_body = doc.add_paragraph()
            # set_urdu_paragraph(p_body, justify=True)
            # p_body.paragraph_format.space_after = Pt(10)
            # p_body.paragraph_format.line_spacing = 1.35
            # run_body = p_body.add_run(sec.get("explanation", ""))
            # set_urdu_run(run_body, font_name='Calibri', size_pt=14, bold=False)

        # Diagram Image
        img_path = image_map.get(q_idx) if isinstance(image_map, dict) else None
        if img_path and os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(10)
            p_img.paragraph_format.space_after = Pt(10)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))

        # Conclusion
        if q.get("conclusion"):
            p_conc_h = doc.add_paragraph()
            set_urdu_paragraph(p_conc_h, justify=True)
            p_conc_h.paragraph_format.space_before = Pt(12)
            p_conc_h.paragraph_format.space_after = Pt(4)
            p_conc_h.paragraph_format.keep_with_next = True
            
            run_conc_h = p_conc_h.add_run("نتیجہ")
            set_urdu_run(run_conc_h, font_name='Calibri', size_pt=15, bold=True,color_rgb=RGBColor(0x00, 0x33, 0x66))

            conclusion_text = q.get("conclusion", "")
            paragraphs_list = [p.strip() for p in conclusion_text.replace("\\n", "\n").split("\n") if p.strip()]

            for para_text in paragraphs_list:
                p_conc_body = doc.add_paragraph()
                set_urdu_paragraph(p_conc_body, justify=True)
                p_conc_body.paragraph_format.space_after = Pt(14)
                p_conc_body.paragraph_format.line_spacing = 1.35
                run_conc_body = p_conc_body.add_run(para_text)
                set_urdu_run(run_conc_body, font_name='Calibri', size_pt=14, bold=False)
            
            # p_conc_body = doc.add_paragraph()
            # set_urdu_paragraph(p_conc_body, justify=True)
            # p_conc_body.paragraph_format.space_after = Pt(14)
            # p_conc_body.paragraph_format.line_spacing = 1.35
            # run_conc_body = p_conc_body.add_run(q.get("conclusion"))
            # set_urdu_run(run_conc_body, font_name='Calibri', size_pt=14, bold=False)

    doc.save(output_path)
    # shutil.rmtree("diagrams", ignore_errors=True)
    return output_path